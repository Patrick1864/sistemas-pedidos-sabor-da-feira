import streamlit as st
import pandas as pd
from datetime import datetime
from docx import Document
import os

# Inicializar estados
if 'pedidos' not in st.session_state:
    st.session_state['pedidos'] = pd.DataFrame(columns=[
        "Data", "Cliente", "Endereço", "Produtos", "Quantidades"
    ])

if 'pedido_editando' not in st.session_state:
    st.session_state['pedido_editando'] = None

# Função para gerar uma ficha Word individual
def gerar_ficha_pedido(pedido, indice):
    doc = Document()
    doc.add_heading(f'Ficha do Pedido - Cliente {pedido["Cliente"]}', level=1)
    doc.add_paragraph(f'Data: {pedido["Data"]}')
    doc.add_paragraph(f'Cliente: {pedido["Cliente"]}')
    doc.add_paragraph(f'Endereço: {pedido["Endereço"]}')
    doc.add_paragraph(f'Produtos: {pedido["Produtos"]}')
    doc.add_paragraph(f'Quantidades: {pedido["Quantidades"]}')
    nome_arquivo = f"pedido_{indice}.docx"
    doc.save(nome_arquivo)
    return nome_arquivo

# Título
st.title("📦 Sistema de Pedidos - Sabor da Feira")

# Formulário de cadastro / edição
st.header("📝 Cadastro / Edição de Pedido")

if st.session_state['pedido_editando'] is not None:
    pedido = st.session_state['pedidos'].iloc[st.session_state['pedido_editando']]
    nome = st.text_input("👤 Nome do cliente:", value=pedido["Cliente"])
    endereco = st.text_input("🏠 Endereço:", value=pedido["Endereço"])
    produtos = st.text_input("📦 Produtos:", value=pedido["Produtos"])
    quantidades = st.text_input("🔢 Quantidades:", value=pedido["Quantidades"])
else:
    nome = st.text_input("👤 Nome do cliente:")
    endereco = st.text_input("🏠 Endereço:")
    produtos = st.text_input("📦 Produtos (separados por vírgula):")
    quantidades = st.text_input("🔢 Quantidades (mesma ordem):")

if st.button("✅ Salvar Pedido"):
    if nome and produtos and quantidades:
        try:
            lista_produtos = [p.strip() for p in produtos.split(",")]
            lista_quantidades = [q.strip() for q in quantidades.split(",")]

            if len(lista_produtos) != len(lista_quantidades):
                st.error("⚠️ Número de produtos e quantidades não corresponde.")
            else:
                novo_pedido = {
                    "Data": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "Cliente": nome,
                    "Endereço": endereco,
                    "Produtos": produtos,
                    "Quantidades": quantidades
                }

                if st.session_state['pedido_editando'] is not None:
                    st.session_state['pedidos'].iloc[st.session_state['pedido_editando']] = novo_pedido
                    st.success("✏️ Pedido atualizado com sucesso!")
                    st.session_state['pedido_editando'] = None
                else:
                    st.session_state['pedidos'] = pd.concat([
                        st.session_state['pedidos'], pd.DataFrame([novo_pedido])
                    ], ignore_index=True)
                    st.success("✅ Pedido cadastrado com sucesso!")
        except Exception as e:
            st.error(f"Erro ao salvar pedido: {e}")
    else:
        st.warning("Preencha todos os campos obrigatórios.")

# 🔍 Busca por nome
st.subheader("🔍 Buscar Pedido por Nome")
nome_busca = st.text_input("Digite o nome do cliente para buscar:")
resultado_busca = st.session_state['pedidos']

if nome_busca:
    resultado_busca = st.session_state['pedidos'][st.session_state['pedidos']['Cliente'].str.contains(nome_busca, case=False, na=False)]

# 📑 Histórico
st.subheader("📑 Histórico de Pedidos")

if resultado_busca.empty:
    st.info("Nenhum pedido encontrado.")
else:
    st.dataframe(resultado_busca, use_container_width=True)

    st.markdown("### ✏️ Editar / Excluir / Baixar Ficha")
    index_para_editar = st.selectbox("Selecione o índice:", options=resultado_busca.index.tolist())

    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("✏️ Editar"):
            st.session_state['pedido_editando'] = index_para_editar
            st.experimental_rerun()

    with col2:
        if st.button("🗑️ Excluir"):
            st.session_state['pedidos'] = st.session_state['pedidos'].drop(index_para_editar).reset_index(drop=True)
            st.success("🗑️ Pedido excluído com sucesso.")
            st.experimental_rerun()

    with col3:
        if st.button("⬇️ Baixar Ficha Word"):
            pedido = st.session_state['pedidos'].loc[index_para_editar]
            caminho = gerar_ficha_pedido(pedido, index_para_editar)
            with open(caminho, "rb") as file:
                st.download_button(
                    label="📄 Baixar Ficha Word",
                    data=file,
                    file_name=caminho,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# 📤 Exportações em lote
st.subheader("📤 Exportar Pedidos")

col1, col2, col3 = st.columns(3)

with col1:
    if st.button("💾 Exportar Excel"):
        try:
            st.session_state['pedidos'].to_excel("pedidos.xlsx", index=False, engine='openpyxl')
            st.success("📁 'pedidos.xlsx' salvo!")
        except Exception as e:
            st.error(f"Erro ao salvar Excel: {e}")

with col2:
    if st.button("📄 Gerar fichas Word individuais"):
        try:
            for i, pedido in st.session_state['pedidos'].iterrows():
                gerar_ficha_pedido(pedido, i)
            st.success("📄 Fichas geradas com sucesso!")
        except Exception as e:
            st.error(f"Erro ao gerar fichas: {e}")

with col3:
    if st.button("📘 Exportar Todos em Word Único"):
        try:
            doc_final = Document()
            doc_final.add_heading('📘 Todos os Pedidos - Sabor da Feira', 0)
            for idx, pedido in st.session_state['pedidos'].iterrows():
                doc_final.add_paragraph("---------------------------")
                doc_final.add_paragraph(f"Pedido #{idx + 1}")
                doc_final.add_paragraph(f"📅 Data: {pedido['Data']}")
                doc_final.add_paragraph(f"👤 Cliente: {pedido['Cliente']}")
                doc_final.add_paragraph(f"🏠 Endereço: {pedido['Endereço']}")
                doc_final.add_paragraph(f"📦 Produtos: {pedido['Produtos']}")
                doc_final.add_paragraph(f"🔢 Quantidades: {pedido['Quantidades']}")

            nome_arquivo = "todos_os_pedidos.docx"
            doc_final.save(nome_arquivo)

            with open(nome_arquivo, "rb") as f:
                st.download_button(
                    label="⬇️ Baixar Todos os Pedidos em Word",
                    data=f,
                    file_name=nome_arquivo,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        except Exception as e:
            st.error(f"Erro ao gerar arquivo Word: {e}")
